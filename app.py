from flask import Flask, request, render_template, jsonify, send_file
import os
import cv2
from ultralytics import YOLO
from transformers import WhisperProcessor, WhisperForConditionalGeneration
import torch
import librosa
from docx import Document
from docx.shared import Inches
from moviepy.editor import VideoFileClip

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['RESULT_FOLDER'] = 'results/'

# Ensure the upload and result directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULT_FOLDER'], exist_ok=True)

# Load your custom YOLO model for vest detection
model = YOLO("best.pt")  # Replace with the path to your trained model

# Load fine-tuned Whisper model and processor
whisper_processor = WhisperProcessor.from_pretrained("./resultsAudio")
whisper_model = WhisperForConditionalGeneration.from_pretrained("./resultsAudio")
whisper_model.eval()  # Set the model to evaluation mode

# Device configuration
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
whisper_model.to(device)

# Words to highlight in the text
words_to_highlight = [
    "Risk", "cost", "delay", "started", "finished", "in progress",
    "caution", "situation", "time overrun", "cost overrun"
]

def create_word_doc(results, docFile, title):
    doc = Document()
    doc.add_heading(f'{title}', 0)
    
    for result in results:
        if 'image' in result:
            doc.add_picture(result['image'], width=Inches(5))
            doc.add_paragraph(result['summary_text'])
        elif 'video' in result:
            doc.add_paragraph(f"Video detected, but video export is not included.")
            doc.add_paragraph(result['summary_text'])
        elif 'fileName' in result:
            doc.add_paragraph(result['fileName'])
            paragraph = doc.add_paragraph()
            
            # Initialize text from summary_text
            text = result['summary_text']
            
            # Sort the words by length to prioritize multi-word phrases (e.g., "cost overrun")
            sorted_words = sorted(words_to_highlight, key=len, reverse=True)
            
            # Replace the words/phrases to mark them
            for word in sorted_words:
                text = text.replace(word, f"||{word}||")

            # Split the text on the placeholders "||"
            parts = text.split("||")
            
            for part in parts:
                run = paragraph.add_run(part)
                if part in words_to_highlight:
                    run.bold = True
                    run.underline = True

    # Save the document to the result folder
    word_doc_path = os.path.join(app.config['RESULT_FOLDER'], f'{docFile}.docx')
    doc.save(word_doc_path)
    return word_doc_path

def delete_file(file_path):
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"File {file_path} has been deleted.")
        else:
            print(f"File {file_path} does not exist.")
    except Exception as e:
        print(f"Error occurred while deleting the file: {e}")

# Highlight specific words in text
def highlight_words(text, words_to_highlight):
    for word in words_to_highlight:
        text = text.replace(word, f'<span class="highlight">{word}</span>')
    return text

def load_audio(file_path, sampling_rate=16000):
    audio, sr = librosa.load(file_path, sr=sampling_rate)
    return audio

# Function to transcribe audio
def transcribe_audio(file_path):
    audio = load_audio(file_path)
    inputs = whisper_processor(audio, sampling_rate=16000, return_tensors="pt").input_features.to(device)

    # Generate transcription
    with torch.no_grad():
        generated_ids = whisper_model.generate(inputs)
    transcription = whisper_processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
    return transcription

# Enhanced vest detection function for images
def detect_vests(image_path):
    img = cv2.imread(image_path)
    result_image_path = os.path.join(app.config['RESULT_FOLDER'], 'result_' + os.path.basename(image_path))
    
    # Perform inference
    results = model(img)
    
    total_people = 0
    people_with_vests = 0
    
    # Process the results
    for result in results:
        boxes = result.boxes.xyxy
        confidences = result.boxes.conf
        class_ids = result.boxes.cls

        for bbox, conf, cls in zip(boxes, confidences, class_ids):
            x1, y1, x2, y2 = [int(coord) for coord in bbox]
            confidence = conf.item()
            cls_id = int(cls.item())
            label = model.names[cls_id]

            total_people += 1
            if label == "vest" and confidence > 0.5:
                people_with_vests += 1
            elif label == "no-vest" and confidence > 0.5:
                cv2.rectangle(img, (x1, y1), (x2, y2), (0, 0, 255), 2)
    
    # Save the annotated image
    cv2.imwrite(result_image_path, img)
    
    # Generate summary text
    people_without_vests = total_people - people_with_vests
    text = 'are'

    if people_without_vests == 1:
        text = "is"
    if people_without_vests == 0:
        summary_text = "Pass, No potentioal risk detected, everyone is following the safety instructions."
    else:
        summary_text = f"Risk of injury is detected, {people_without_vests} person {text} not following the safety instructions, an action is needed."
    
    return result_image_path, summary_text

# Enhanced vest detection function for videos
def process_video(video_path):
    result_video_path = os.path.join(app.config['RESULT_FOLDER'], 'result_' + os.path.basename(video_path))
    cap = cv2.VideoCapture(video_path)
    fourcc = cv2.VideoWriter_fourcc(*'mp4v')
    out = cv2.VideoWriter(result_video_path, fourcc, 20.0, (int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)), int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))))
    
    total_people = 0
    people_with_vests = 0
    
    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break
        
        # Perform inference
        results = model(frame)
        
        for result in results:
            boxes = result.boxes.xyxy
            confidences = result.boxes.conf
            class_ids = result.boxes.cls

            for bbox, conf, cls in zip(boxes, confidences, class_ids):
                x1, y1, x2, y2 = [int(coord) for coord in bbox]
                confidence = conf.item()
                cls_id = int(cls.item())
                label = model.names[cls_id]

                total_people += 1
                if label == "vest" and confidence > 0.5:
                    people_with_vests += 1
                    cv2.rectangle(frame, (x1, y1), (x2, y2), (0, 255, 0), 2)
                    cv2.putText(frame, f"{label} {confidence:.2f}", (x1, y1 - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (0, 255, 0), 2)
                elif label == "no-vest" and confidence > 0.5:
                    cv2.rectangle(frame, (x1, y1), (x2, y2), (0, 0, 255), 2)
        
        out.write(frame)
    
    cap.release()
    out.release()
    
    people_without_vests = total_people - people_with_vests
    if people_without_vests == 0:
        summary_text = "Pass, everyone is following the safety instructions."
    else:
        summary_text = f"Action needed, {people_without_vests} people are not following the safety instructions."
    
    return result_video_path, summary_text

# Function to extract audio from video
def extract_audio_from_video(video_path):
    # Ensure the file exists
    if not os.path.exists(video_path):
        print(f"Error: The file {video_path} does not exist.")
        return None

    # Use os.path.join to construct the audio path
    audio_filename = 'audio_' + os.path.splitext(os.path.basename(video_path))[0] + '.wav'
    audio_path = os.path.join(app.config['UPLOAD_FOLDER'], audio_filename)

    try:
        # Extract audio using moviepy
        video = VideoFileClip(video_path)
        video.audio.write_audiofile(audio_path)
        print(f"Audio successfully extracted to: {audio_path}")
    except Exception as e:
        print(f"Error extracting audio from video: {e}")
        return None

    return audio_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload_image', methods=['GET', 'POST'])
def upload_image():
    if request.method == 'POST':
        files = request.files.getlist('file')
        results = []
        for file in files:
            if file:
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(filepath)
                if file.filename.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
                    result_video_path, summary_text = process_video(filepath)
                    results.append({'video': result_video_path, 'summary_text': summary_text})
                    delete_file(filepath)
                else:
                    result_image_path, summary_text = detect_vests(filepath)
                    results.append({'image': result_image_path, 'summary_text': summary_text})
                    delete_file(filepath)
        create_word_doc(results,'vest_detection_results','Vest Detection Results')
        return jsonify(results)
       # return render_template('result.html', results=results, title="Image/Video Results")
    return render_template('detect_vest.html')

@app.route('/upload_audio', methods=['GET', 'POST'])
def upload_audio():
    if request.method == 'POST':
        files = request.files.getlist('file')
        audio_results = []
        origenText = []
        for file in files:
            if file:
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(filepath)
                text = transcribe_audio(filepath)
                origenText.append({'fileName': file.filename, 'summary_text': text})
                highlighted_text = highlight_words(text, words_to_highlight)
                audio_results.append({'fileName': file.filename, 'summary_text': highlighted_text})
                create_word_doc(origenText,'audio_detection_results','Audio Detection Results')
                delete_file(filepath)
                

        return render_template('audio_results.html', results=audio_results, title="Audio to Text Results")
    return render_template('audio_to_text.html')

@app.route('/upload_video', methods=['GET', 'POST'])
def upload_video():
    if request.method == 'POST':
        file = request.files['file']
        video_results = []
        origenText = []
        if file:
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filepath)
            # Reuse the transcribe_audio function to extract and transcribe audio from video
            audio_path = extract_audio_from_video(filepath)
            text = transcribe_audio(audio_path)
            origenText.append({'fileName': file.filename, 'summary_text': text})
            highlighted_text = highlight_words(text, words_to_highlight)
            video_results.append({'fileName': file.filename, 'summary_text': highlighted_text})
            create_word_doc(origenText,'video_detection_results','Video Detection Results')
            delete_file(filepath)
            delete_file(audio_path)
            return render_template('audio_results.html', results=video_results, title="Video to Text Results")
    return render_template('video_to_text.html')

@app.route('/display_image/<path:image_path>')
def display_image(image_path):
    return send_file(image_path, mimetype='image/jpeg')

@app.route('/display_video/<path:video_path>')
def display_video(video_path):
    return send_file(video_path, mimetype='video/mp4')

if __name__ == '__main__':
    app.run(debug=True)
